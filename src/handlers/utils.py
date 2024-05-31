
import io
from datetime import datetime, timedelta
from typing import Dict

import requests
from aiogram.types.input_file import BufferedInputFile
from docx import Document
# from docx2pdf import convert
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Inches, Pt

from src.handlers.schemas import FacturaLimitInfo, OrderResponse, ProductPrices


def format_stir(phone_number):
    formatted_number = ' '.join([phone_number[i:i+3]
                                 for i in range(0, len(phone_number), 3)])
    return formatted_number


def format_number(number):
    """
    Formats the given number for human readability with space separator.

    Args:
        number (int or float): The number to format.

    Returns:
        str: The formatted number.
    """
    formatted_number = f"{number:,}".replace(',', ' ')
    part = formatted_number.split('.')
    if len(part) == 2 and part[1] == '0':
        return part[0]
    return formatted_number


def format_phone(phone_number: str):
    if "+" in phone_number:
        phone_number = phone_number.replace('+', '')
    formatted_number = f"+{phone_number[0:3]} ({phone_number[3:5]}) {phone_number[5:8]}-{phone_number[8:10]}-{phone_number[10:12]}"
    return formatted_number


def create_facture(data: OrderResponse, prices: Dict) -> BufferedInputFile:
    """
        hisob faktura yaratadi
    """
    # Создание документа
    doc = Document()
    order_id = data.sequence_number
    # Настройка шрифта по умолчанию
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    para_format = style.paragraph_format
    #
    doc.sections[0].bottom_margin = Inches(0.4)
    doc.sections[0].left_margin = Inches(0.4)
    doc.sections[0].top_margin = Inches(0.4)
    doc.sections[0].right_margin = Inches(0.4)
    # header qismini qo'shsih
    header = doc.add_paragraph()
    header.add_run("___.05.2024 sanadagi _____-sonli shartnomaga")
    header.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    header.paragraph_format.space_after = Pt(0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header2 = doc.add_paragraph()
    current_date = str(get_current_date())
    header2.add_run(f"{current_date} sanadagi {order_id}-sonli")
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header2.paragraph_format.space_after = Pt(0)
    header2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # header.paragraph_format.line_spacing = Pt(1.0)

    # ===================================
    empty_pg = doc.add_paragraph()
    run = empty_pg.add_run("")
    run.font.size = Pt(20)
    empty_pg.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    empty_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # ==============================================
    header = doc.add_paragraph()
    run = header.add_run("YUK XATI")
    run.font.size = Pt(26)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -------------------------------------

    # Добавление информации о поставщике и получателе

    yuridik_data = [
        [
            {"Yetkazib beruvchi: ": data.company.name},
            {"Qabul qiluvchi: ": data.dmtt.name},
        ],
        [
            {"Manzil: ": data.company.address},
            {"Manzil: ": data.dmtt.address},
        ],
        [{"Tel: ": format_phone(data.company.phone_number)},
         {"Tel: ": format_phone(data.dmtt.user.phone_number)}],
        [{"STIR: ": format_stir(data.company.stir)}, {
            "STIR: ": format_stir(data.dmtt.stir)}],
    ]

    table = doc.add_table(rows=4, cols=2)

    for i in range(4):
        for j in range(2):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            item = yuridik_data[i][j]
            key, value = item.popitem()
            run = p.add_run(key)
            run.bold = True
            p.add_run(value)

            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # ------------------------------------------------------------
    for _ in range(2):
        p = doc.add_paragraph()
        run = p.add_run()
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
    # ----------------------
    # Добавление таблицы с товарами
    products_table = doc.add_table(rows=1, cols=7)
    products_table.style = "Table Grid"
    columns_width = {
        0: 0.42,
        1: 2.31,
        2: 0.95,
        3: 0.92,
        4: 0.98,
        5: 0.89,
        6: 1.38
    }
    headers = [
        "T/r",
        "Mahsulot nomi",
        "O'lchov birligi",
        "Miqdori",
        "Narxi (so'm)",
        "QQS",
        "Yetkazib berish qiymati",
    ]
    for i, header in enumerate(headers):
        cell = products_table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True  # Sarlavhani qalin qilish
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Заполнение таблицы данными
    i = 0
    total_summ = 0
    for item in data.items:
        price_item = prices.get(item.product_name)
        if not price_item:
            price_item = {
                "measure": "kg",
                "price": 0
            }
        i += 1
        row_cells = products_table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = item.product_name  # Примерное название
        row_cells[2].text = price_item.get('measure')
        row_cells[3].text = format_number(item.count)  # Примерное количество
        row_cells[4].text = format_number(price_item.get('price'))
        row_cells[5].text = "QQS siz"
        summa = item.count*int(price_item.get('price'))
        row_cells[6].text = format_number(summa)
        total_summ += summa

        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for cell in row_cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)
            cell.paragraphs[0].paragraph_format.space_before = Pt(3)

    # set size
    for column_idx in range(len(products_table.columns)):
        for j, cell in enumerate(products_table.columns[column_idx].cells):
            cell.width = Inches(columns_width[column_idx])
    # ------------------------

    doc.add_paragraph()

    # summary
    header = doc.add_paragraph()
    run = header.add_run(
        f"Jami yetkazib berilgan mahsulotlarning umumiy qiymati - {format_number(total_summ)} so'm"
    )
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # space = doc.add_paragraph()

    table = doc.add_table(rows=2, cols=2)

    table.cell(0, 0).text = "Yetkazib beruvchi:"
    table.cell(0, 1).text = "Qabul qiluvchi:"
    table.cell(1, 0).text = "M. Abdubannobov"
    # dmtt  zavhoz nomi
    table.cell(
        1, 1).text = f"{data.dmtt.user.first_name} {data.dmtt.user.last_name}"

    # doc.save("doc.docx")
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    dd = BufferedInputFile(
        file_stream.read(), filename=f"№{order_id} ({data.dmtt.name}).docx")
    return dd


def create_facture_without_price(data: OrderResponse, prices: Dict) -> BufferedInputFile:
    """
        hisob faktura yaratadi narxlarsiz
    """
    # Создание документа
    doc = Document()
    order_id = data.sequence_number
    # Настройка шрифта по умолчанию
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    para_format = style.paragraph_format
    #
    doc.sections[0].bottom_margin = Inches(0.4)
    doc.sections[0].left_margin = Inches(0.4)
    doc.sections[0].top_margin = Inches(0.4)
    doc.sections[0].right_margin = Inches(0.4)
    # header qismini qo'shsih
    header = doc.add_paragraph()
    header.add_run("___.05.2024 sanadagi _____-sonli shartnomaga")
    header.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    header.paragraph_format.space_after = Pt(0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header2 = doc.add_paragraph()
    current_date = str(get_current_date())
    header2.add_run(f"{current_date} sanadagi {order_id}-sonli")
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header2.paragraph_format.space_after = Pt(0)
    header2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # header.paragraph_format.line_spacing = Pt(1.0)

    # ===================================
    empty_pg = doc.add_paragraph()
    run = empty_pg.add_run("")
    run.font.size = Pt(20)
    empty_pg.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    empty_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # ==============================================
    header = doc.add_paragraph()
    run = header.add_run("YUK XATI")
    run.font.size = Pt(26)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -------------------------------------

    # Добавление информации о поставщике и получателе

    yuridik_data = [
        [
            {"Yetkazib beruvchi: ": data.company.name},
            {"Qabul qiluvchi: ": data.dmtt.name},
        ],
        [
            {"Manzil: ": data.company.address},
            {"Manzil: ": data.dmtt.address},
        ],
        [{"Tel: ": format_phone(data.company.phone_number)},
         {"Tel: ": format_phone(data.dmtt.user.phone_number)}],
        [{"STIR: ": format_stir(data.company.stir)}, {
            "STIR: ": format_stir(data.dmtt.stir)}],
    ]

    table = doc.add_table(rows=4, cols=2)

    for i in range(4):
        for j in range(2):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            item = yuridik_data[i][j]
            key, value = item.popitem()
            run = p.add_run(key)
            run.bold = True
            p.add_run(value)

            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # ------------------------------------------------------------
    for _ in range(2):
        p = doc.add_paragraph()
        run = p.add_run()
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
    # ----------------------
    # Добавление таблицы с товарами
    products_table = doc.add_table(rows=1, cols=7)
    products_table.style = "Table Grid"
    columns_width = {
        0: 0.42,
        1: 2.31,
        2: 0.95,
        3: 0.92,
        4: 0.98,
        5: 0.89,
        6: 1.38
    }
    headers = [
        "T/r",
        "Mahsulot nomi",
        "O'lchov birligi",
        "Miqdori",
        "Narxi (so'm)",
        "QQS",
        "Yetkazib berish qiymati",
    ]
    for i, header in enumerate(headers):
        cell = products_table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True  # Sarlavhani qalin qilish
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Заполнение таблицы данными
    i = 0
    total_summ = 0
    for item in data.items:
        price_item = prices.get(item.product_name)
        if not price_item:
            price_item = {
                "measure": "kg",
            }
        i += 1
        row_cells = products_table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = item.product_name  # Примерное название
        row_cells[3].text = format_number(item.count)  # Примерное количество
        row_cells[4].text = ""
        row_cells[2].text = price_item.get('measure')
        row_cells[5].text = "QQS siz"
        row_cells[6].text = ""

        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for cell in row_cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)
            cell.paragraphs[0].paragraph_format.space_before = Pt(3)

    # set size
    for column_idx in range(len(products_table.columns)):
        for j, cell in enumerate(products_table.columns[column_idx].cells):
            cell.width = Inches(columns_width[column_idx])
    # ------------------------

    doc.add_paragraph()

    # summary
    header = doc.add_paragraph()
    run = header.add_run(
        f"Jami yetkazib berilgan mahsulotlarning umumiy qiymati - so'm"
    )
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # space = doc.add_paragraph()

    table = doc.add_table(rows=2, cols=2)

    table.cell(0, 0).text = "Yetkazib beruvchi:"
    table.cell(0, 1).text = "Qabul qiluvchi:"
    table.cell(1, 0).text = "M. Abdubannobov"
    # dmtt  zavhoz nomi
    table.cell(
        1, 1).text = f"{data.dmtt.user.first_name} {data.dmtt.user.last_name}"

    # doc.save("doc.docx")
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    dd = BufferedInputFile(
        file_stream.read(), filename=f"№{order_id} ({data.dmtt.name}).docx")
    return dd


def get_current_date():
    # Get the current date
    current_date = datetime.today()

    # Format the date
    formatted_date = current_date.strftime("%d.%m.%Y")

    # Print the formatted date
    return formatted_date


def get_order_as_list(data, order_id):
    """
        retun order as list
    """
    malumot = f"<b>{data['dmtt']['name']} | Buyurtma - {order_id}</b>\n\n"
    index = 0
    for i in data["items"]:
        index += 1
        count = format_number(i['count'])
        malumot += f"{index}. {i['product_name']} - {count}\n"
    return malumot


def get_first_and_last_day_of_current_month():
    # Hozirgi sanani olish
    today = datetime.today()

    # Oyning birinchi kunini olish
    first_day = today.replace(day=1)

    # Oyning oxirgi kunini olish
    # Kelgusi oydan bir kun oldin
    next_month = first_day.replace(day=28) + timedelta(days=4)
    last_day = next_month - timedelta(days=next_month.day)

    # Sanalarni formatlash
    first_day_str = first_day.strftime("%d.%m")
    last_day_str = last_day.strftime("%d.%m.%Y")

    return f"{first_day_str}-{last_day_str}"


def create_full_facture(order_id: int, data: FacturaLimitInfo, prices: Dict) -> BufferedInputFile:
    """
        hisob faktura yaratadi
    """
    # Создание документа
    doc = Document()

    # Настройка шрифта по умолчанию
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    #
    doc.sections[0].bottom_margin = Inches(0.4)
    doc.sections[0].left_margin = Inches(0.4)
    doc.sections[0].top_margin = Inches(0.4)
    doc.sections[0].right_margin = Inches(0.4)
    # header qismini qo'shsih
    header = doc.add_paragraph()
    header.add_run("___.05.2024 sanadagi _____-sonli shartnomaga")
    header.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    header.paragraph_format.space_after = Pt(0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header2 = doc.add_paragraph()
    header2.add_run(
        f"{get_first_and_last_day_of_current_month()} sanalardagi {order_id}-sonli")
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header2.paragraph_format.space_after = Pt(0)
    header2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # header.paragraph_format.line_spacing = Pt(1.0)

    # ===================================
    empty_pg = doc.add_paragraph()
    run = empty_pg.add_run("")
    run.font.size = Pt(20)
    empty_pg.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    empty_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # ==============================================
    header = doc.add_paragraph()
    run = header.add_run("HISOBVARAQ FAKTURA")
    run.font.size = Pt(26)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -------------------------------------

    # Добавление информации о поставщике и получателе

    yuridik_data = [
        [
            {"Yetkazib beruvchi: ": data.company.name},
            {"Qabul qiluvchi: ": data.dmtt.name},
        ],
        [
            {"Manzil: ": data.company.address},
            {"Manzil: ": data.dmtt.address},
        ],
        [{"Tel: ": format_phone(data.company.phone_number)},
         {"Tel: ": format_phone(data.dmtt.user.phone_number)}],
        [{"STIR: ": format_stir(data.company.stir)}, {
            "STIR: ": format_stir(data.dmtt.stir)}],
    ]

    table = doc.add_table(rows=4, cols=2)

    for i in range(4):
        for j in range(2):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            item = yuridik_data[i][j]
            key, value = item.popitem()
            run = p.add_run(key)
            run.bold = True
            p.add_run(value)

            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # ------------------------------------------------------------
    for _ in range(2):
        p = doc.add_paragraph()
        run = p.add_run()
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
    # ----------------------
    # Добавление таблицы с товарами
    products_table = doc.add_table(rows=1, cols=7)
    products_table.style = "Table Grid"
    columns_width = {
        0: 0.42,
        1: 2.31,
        2: 0.95,
        3: 0.92,
        4: 0.98,
        5: 0.89,
        6: 1.38
    }

    headers = [
        "T/r",
        "Mahsulot nomi",
        "O'lchov birligi",
        "Miqdori",
        "Narxi (so'm)",
        "QQS",
        "Yetkazib berish qiymati",
    ]
    for i, header in enumerate(headers):
        cell = products_table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True  # Sarlavhani qalin qilish
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Заполнение таблицы данными
    i = 0
    total_summ = 0
    for item in data.items:
        price_item = prices.get(item.name)
        if not price_item:
            price_item = {
                "measure": "kg",
                "price": 0
            }
        i += 1
        row_cells = products_table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = item.name  # Примерное название
        row_cells[2].text = price_item.get('measure')
        row_cells[3].text = format_number(
            float(item.limit))  # Примерное количество
        row_cells[4].text = format_number(price_item.get('price'))
        row_cells[5].text = "QQS siz"
        summa = float(item.limit)*int(price_item.get('price'))
        row_cells[6].text = format_number(summa)
        total_summ += summa

        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for cell in row_cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)
            cell.paragraphs[0].paragraph_format.space_before = Pt(3)

    # set size
    for column_idx in range(len(products_table.columns)):
        for j, cell in enumerate(products_table.columns[column_idx].cells):
            cell.width = Inches(columns_width[column_idx])
    # ------------------------

    doc.add_paragraph()

    # summary
    header = doc.add_paragraph()
    run = header.add_run(
        f"Jami yetkazib berilgan mahsulotlarning umumiy qiymati - {format_number(total_summ)} so'm"
    )
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # space = doc.add_paragraph()

    table = doc.add_table(rows=2, cols=2)

    table.cell(0, 0).text = "Yetkazib beruvchi:"
    table.cell(0, 1).text = "Qabul qiluvchi:"
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table.cell(1, 0).text = "M. Abdubannobov"
    # dmtt  zavhoz nomi
    table.cell(
        1, 1).text = f"{data.dmtt.user.first_name} {data.dmtt.user.last_name}"
    table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # doc.save("doc.docx")
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    dd = BufferedInputFile(
        file_stream.read(), filename=f"№{order_id} ({data.dmtt.name}).docx")
    return dd
